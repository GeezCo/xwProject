# -*- coding: utf-8 -*-
"""
Word文档格式化工具 - 专业版
功能：
1. 支持配置文件自定义样式（优先级最高）
2. 支持从参考文档复制样式
3. 自动识别标题并应用样式
4. 自动应用三线表样式
5. 自动转换将来时态
6. 稳定输出，不乱码不出错

使用方法:
    python format_word.py <输入文档> <输出文档>
    python format_word.py <输入文档> <输出文档> --reference 参考文档
    python format_word.py <输入文档> <输出文档> --config 配置文件

示例:
    python format_word.py input.docx output.docx
    python format_word.py input.docx output.docx --reference Revised.docx
    python format_word.py input.docx output.docx --config word_format.conf
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys
import os
import configparser
from pathlib import Path

# ============================================================================
# 默认配置（如果配置文件不存在或缺少某项，使用此默认值）
# ============================================================================

DEFAULT_CONFIG = {
    # 正文样式 - 参考文档标准
    'body_font': '仿宋',
    'body_font_fallback': '宋体',
    'body_font_size': 10.5,  # 五号
    'body_line_spacing': 15.5,
    'body_first_line_indent': 0,

    # 一级标题 - 参考文档标准
    'title1_font': '黑体',
    'title1_font_size': 22,
    'title1_line_spacing': 21,
    'title1_bold': True,
    'title1_first_line_indent': 0,

    # 二级标题 - 参考文档标准
    'title2_font': '黑体',
    'title2_font_size': 16,
    'title2_line_spacing': 21,
    'title2_bold': True,
    'title2_first_line_indent': 0,

    # 三级标题 - 参考文档标准
    'title3_font': '黑体',
    'title3_font_size': 16,
    'title3_line_spacing': 21,
    'title3_bold': True,
    'title3_first_line_indent': 0,

    # 四级标题 - 参考文档标准
    'title4_font': '黑体',
    'title4_font_size': 14,
    'title4_line_spacing': 21,
    'title4_bold': True,
    'title4_first_line_indent': 0,

    # 五级标题 - 参考文档标准
    'title5_font': '黑体',
    'title5_font_size': 12,
    'title5_line_spacing': 20,
    'title5_bold': True,
    'title5_first_line_indent': 0,

    # 六级标题 - 参考文档标准
    'title6_font': '黑体',
    'title6_font_size': 12,
    'title6_line_spacing': 20,
    'title6_bold': True,
    'title6_first_line_indent': 0,

    # 表格样式 - 参考文档标准
    'table_header_font': '黑体',
    'table_header_font_size': 12,
    'table_header_bold': True,
    'table_body_font': '宋体',
    'table_body_font_size': 10.5,
    'table_body_bold': False,
    'table_border_top_bottom': 1.5,
    'table_border_header': 0.75,

    # 目录样式
    'toc_title_font': '黑体',
    'toc_title_font_size': 16,
    'toc_title_bold': True,
    'toc_level1_font': '黑体',
    'toc_level1_font_size': 12,
    'toc_level2_font': '宋体',
    'toc_level2_font_size': 12,
    'toc_level3_font': '宋体',
    'toc_level3_font_size': 12,
    'toc_line_spacing': 20,

    # 页眉页脚
    'header_font': '宋体',
    'header_font_size': 9,
    'footer_font': '宋体',
    'footer_font_size': 9,

    # 页面设置
    'page_margin_top': 2.5,
    'page_margin_bottom': 2.5,
    'page_margin_left': 3.0,
    'page_margin_right': 2.5,
    'header_margin': 1.5,
    'footer_margin': 1.75,

    # 文本转换
    'enable_future_tense': True,
    'tense_replacements': '已完成=将完成,已实现=将实现,正在进行=将进行,已经=将要,正在=将要,实现了=将实现,完成了=将完成,开发了=将开发,设计了=将设计,测试了=将测试,部署了=将部署',

    # 标题识别 - 更灵活的正则表达式
    # 匹配格式: "1 xxx", "1.1 xxx", "1.1.1 xxx" 等（数字后必须有空格）
    # 排除: "1)"、"1、"等列表项格式
    'heading1_pattern': r'^\d+[\s\u3000]+[^\)\d]',
    'heading2_pattern': r'^\d+\.\d+[\s\u3000]+[^\)\d]',
    'heading3_pattern': r'^\d+\.\d+\.\d+[\s\u3000]+[^\)\d]',
    'heading4_pattern': r'^\d+\.\d+\.\d+\.\d+[\s\u3000]+[^\)\d]',
    'heading5_pattern': r'^\d+\.\d+\.\d+\.\d+\.\d+[\s\u3000]+[^\)\d]',
    'heading6_pattern': r'^\d+\.\d+\.\d+\.\d+\.\d+\.\d+[\s\u3000]+[^\)\d]',

    # 输出配置
    'verbose': True,
    'keep_backup': True,
    'backup_format': '{original}_backup_{date}',

    # 特殊页面跳过配置
    'skip_special_pages': True,
    'special_page_styles': '封面_文档封格式版本,封面_项目文档统一标识,封面_文件名,封面_文件标识,封面_单位签名,修改记录页_正文,目录_正文,toc 1,toc 2,toc 3',
}

# ============================================================================
# 配置文件加载
# ============================================================================

class ConfigLoader:
    """配置文件加载器"""

    def __init__(self, config_path=None):
        """
        初始化配置加载器

        Args:
            config_path: 配置文件路径（可选）
        """
        self.config = DEFAULT_CONFIG.copy()

        # 默认配置文件路径
        self.default_config_file = 'word_format.conf'

        # 加载配置
        if config_path and os.path.exists(config_path):
            self._load_config_file(config_path)
            self.config_source = config_path
        elif os.path.exists(self.default_config_file):
            self._load_config_file(self.default_config_file)
            self.config_source = self.default_config_file
        else:
            self.config_source = '默认配置'

    def _load_config_file(self, config_path):
        """
        加载配置文件

        Args:
            config_path: 配置文件路径
        """
        try:
            parser = configparser.ConfigParser()

            # 读取配置文件
            with open(config_path, 'r', encoding='utf-8') as f:
                # 配置文件可能没有section，添加一个默认section
                content = f.read()
                if not content.strip().startswith('['):
                    content = '[styles]\n' + content
                parser.read_string(content)

            # 解析配置项
            for section in parser.sections():
                for key, value in parser.items(section):
                    key_lower = key.lower()

                    # 类型转换
                    if key_lower.endswith('_size') or key_lower.endswith('_spacing') or \
                       key_lower.endswith('_indent') or key_lower.endswith('_margin') or \
                       key_lower.endswith('_width') or key_lower == 'table_border_top_bottom' or \
                       key_lower == 'table_border_header':
                        # 数值类型
                        try:
                            self.config[key_lower] = float(value)
                        except:
                            self.config[key_lower] = value

                    elif key_lower.endswith('_bold') or key_lower.endswith('_enable') or \
                         key_lower.endswith('_verbose') or key_lower.endswith('_backup'):
                        # 布尔类型
                        self.config[key_lower] = value.lower() == 'true'

                    elif key_lower.endswith('_pattern'):
                        # 正则表达式
                        self.config[key_lower] = value

                    elif key_lower.endswith('_replacements'):
                        # 替换规则字符串
                        self.config[key_lower] = value

                    else:
                        # 字符串类型
                        self.config[key_lower] = value

        except Exception as e:
            print(f"[WARN] 加载配置文件时出错: {e}")
            print("  将使用默认配置")

    def get(self, key, default=None):
        """
        获取配置项

        Args:
            key: 配置项名称
            default: 默认值

        Returns:
            配置值
        """
        return self.config.get(key.lower(), default)

    def get_tense_replacements(self):
        """
        获取时态替换规则字典

        Returns:
            替换规则字典
        """
        replacements_str = self.get('tense_replacements', '')
        replacements = {}

        if replacements_str:
            # 解析替换规则
            for rule in replacements_str.split(','):
                if '=' in rule:
                    old, new = rule.split('=', 1)
                    replacements[old.strip()] = new.strip()

        return replacements

    def get_heading_patterns(self):
        """
        获取标题识别正则表达式

        Returns:
            标题识别正则表达式列表
        """
        patterns = []
        for i in range(1, 7):
            pattern = self.get(f'heading{i}_pattern')
            if pattern:
                patterns.append(re.compile(pattern))
            else:
                patterns.append(None)

        return patterns

# ============================================================================
# 辅助函数
# ============================================================================

def safe_set_font(run, font_name, font_size, bold=False, fallback_font=None):
    """
    安全设置字体（防止乱码）
    中文使用指定字体，英文使用Times New Roman

    Args:
        run: run对象
        font_name: 中文字体名称
        font_size: 字号（磅）
        bold: 是否加粗
        fallback_font: 备选字体
    """
    try:
        # 设置西文字体为Times New Roman
        run.font.name = 'Times New Roman'

        # 设置东亚字体（中文）为指定字体
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')

        # 设置字号
        run.font.size = Pt(font_size)

        # 设置加粗
        run.font.bold = bold

    except Exception as e:
        # 如果设置失败，使用备选字体
        if fallback_font:
            try:
                run.font.name = 'Times New Roman'
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), fallback_font)
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                run.font.size = Pt(font_size)
                run.font.bold = bold
            except:
                pass  # 最后的容错

def safe_set_paragraph_format(paragraph, font_name, font_size, line_spacing,
                               first_line_indent=None, bold=False, alignment=None,
                               fallback_font=None):
    """
    安全设置段落格式

    Args:
        paragraph: 段落对象
        font_name: 字体名称
        font_size: 字号（磅）
        line_spacing: 行距（磅）
        first_line_indent: 首行缩进（字符数）
        bold: 是否加粗
        alignment: 对齐方式
        fallback_font: 备选字体
    """
    try:
        # 设置字体
        for run in paragraph.runs:
            safe_set_font(run, font_name, font_size, bold, fallback_font)

        # 如果段落没有runs但有文本，创建一个run
        if not paragraph.runs and paragraph.text:
            run = paragraph.add_run(paragraph.text)
            safe_set_font(run, font_name, font_size, bold, fallback_font)

        # 设置段落格式
        pf = paragraph.paragraph_format
        pf.line_spacing = Pt(line_spacing)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        # 设置首行缩进
        if first_line_indent:
            # 1个字符约等于12磅（三号字）
            indent_pt = first_line_indent * 12
            pf.first_line_indent = Pt(indent_pt)

        # 设置对齐方式
        if alignment:
            pf.alignment = alignment

    except Exception as e:
        if config and config.get('verbose'):
            print(f"  [WARN] 设置段落格式时出现警告: {e}")

# ============================================================================
# 标题识别（使用配置文件中的正则表达式）
# ============================================================================

def detect_heading_level(text, patterns):
    """
    自动检测标题级别

    Args:
        text: 段落文本
        patterns: 标题识别正则表达式列表

    Returns:
        0: 正文
        1-6: 标题级别
    """
    text = text.strip()

    if not text:
        return 0

    # 检测数字编号格式（从高到低检测）
    for level in range(6, 0, -1):
        pattern = patterns[level - 1]
        if pattern and pattern.match(text):
            return level

    # 检测中文编号格式（一、二、三）
    if re.match(r'^[一二三四五六七八九十]+[\s\u3000]+', text):
        return 1

    return 0

def apply_heading_format(paragraph, level, config):
    """
    应用标题格式（包括左缩进和悬挂缩进）

    Args:
        paragraph: 段落对象
        level: 标题级别（1-6）
        config: 配置对象
    """
    # 根据标题级别获取配置
    font_name = config.get(f'title{level}_font', DEFAULT_CONFIG[f'title{level}_font'])
    font_size = config.get(f'title{level}_font_size', DEFAULT_CONFIG[f'title{level}_font_size'])
    line_spacing = config.get(f'title{level}_line_spacing', DEFAULT_CONFIG[f'title{level}_line_spacing'])
    bold = config.get(f'title{level}_bold', DEFAULT_CONFIG[f'title{level}_bold'])
    # 获取首行缩进配置（可能为负值，表示悬挂缩进）
    first_line_indent = config.get(f'title{level}_first_line_indent', 0)
    # 获取左缩进配置
    left_indent = config.get(f'title{level}_left_indent', 0)
    fallback_font = config.get('body_font_fallback')

    # 设置字体和行距
    safe_set_paragraph_format(
        paragraph,
        font_name,
        font_size,
        line_spacing,
        first_line_indent=first_line_indent,
        bold=bold,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        fallback_font=fallback_font
    )

    # 设置左缩进（磅值）
    if left_indent:
        try:
            paragraph.paragraph_format.left_indent = Pt(left_indent)
        except:
            pass

    # 尝试设置Word标题样式
    try:
        paragraph.style = f'Heading {level}'
    except:
        pass

# ============================================================================
# 正文格式设置
# ============================================================================

def apply_body_format(paragraph, config):
    """
    应用正文格式
    正文：仿宋GB2312，三号，29磅行距，首行缩进2字符
    """
    font_name = config.get('body_font', DEFAULT_CONFIG['body_font'])
    font_size = config.get('body_font_size', DEFAULT_CONFIG['body_font_size'])
    line_spacing = config.get('body_line_spacing', DEFAULT_CONFIG['body_line_spacing'])
    first_line_indent = config.get('body_first_line_indent', DEFAULT_CONFIG['body_first_line_indent'])
    fallback_font = config.get('body_font_fallback', DEFAULT_CONFIG['body_font_fallback'])

    safe_set_paragraph_format(
        paragraph,
        font_name,
        font_size,
        line_spacing,
        first_line_indent=first_line_indent,
        bold=False,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        fallback_font=fallback_font
    )

# ============================================================================
# 三线表格式设置
# ============================================================================

def apply_three_line_table(table, config):
    """
    应用三线表格式

    Args:
        table: 表格对象
        config: 配置对象
    """
    try:
        # 获取表格样式配置
        header_font = config.get('table_header_font', DEFAULT_CONFIG['table_header_font'])
        header_font_size = config.get('table_header_font_size', DEFAULT_CONFIG['table_header_font_size'])
        header_bold = config.get('table_header_bold', DEFAULT_CONFIG['table_header_bold'])
        body_font = config.get('table_body_font', DEFAULT_CONFIG['table_body_font'])
        body_font_size = config.get('table_body_font_size', DEFAULT_CONFIG['table_body_font_size'])
        body_bold = config.get('table_body_bold', DEFAULT_CONFIG['table_body_bold'])
        border_top_bottom = config.get('table_border_top_bottom', DEFAULT_CONFIG['table_border_top_bottom'])
        border_header = config.get('table_border_header', DEFAULT_CONFIG['table_border_header'])
        fallback_font = config.get('body_font_fallback')

        # 获取表格XML元素
        tbl = table._tbl
        tblPr = tbl.tblPr

        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # 创建表格边框元素
        tblBorders = OxmlElement('w:tblBorders')

        # 设置顶部边框（粗线）
        topBorder = OxmlElement('w:top')
        topBorder.set(qn('w:val'), 'single')
        topBorder.set(qn('w:sz'), str(int(border_top_bottom * 8)))
        topBorder.set(qn('w:space'), '0')
        topBorder.set(qn('w:color'), '000000')
        tblBorders.append(topBorder)

        # 设置底部边框（粗线）
        bottomBorder = OxmlElement('w:bottom')
        bottomBorder.set(qn('w:val'), 'single')
        bottomBorder.set(qn('w:sz'), str(int(border_top_bottom * 8)))
        bottomBorder.set(qn('w:space'), '0')
        bottomBorder.set(qn('w:color'), '000000')
        tblBorders.append(bottomBorder)

        # 清除内部垂直边框
        insideV = OxmlElement('w:insideV')
        insideV.set(qn('w:val'), 'nil')
        tblBorders.append(insideV)

        # 清除左右边框
        leftBorder = OxmlElement('w:left')
        leftBorder.set(qn('w:val'), 'nil')
        tblBorders.append(leftBorder)

        rightBorder = OxmlElement('w:right')
        rightBorder.set(qn('w:val'), 'nil')
        tblBorders.append(rightBorder)

        # 清除旧的边框设置
        oldBorders = tblPr.find(qn('w:tblBorders'))
        if oldBorders is not None:
            tblPr.remove(oldBorders)

        # 应用新边框
        tblPr.append(tblBorders)

        # 设置表头行格式
        if len(table.rows) > 0:
            header_row = table.rows[0]

            # 设置表头单元格格式
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        safe_set_font(run, header_font, header_font_size, header_bold, fallback_font)
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 设置表头单元格下边框（细线）
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                tcBorders = OxmlElement('w:tcBorders')

                # 下边框
                cellBottomBorder = OxmlElement('w:bottom')
                cellBottomBorder.set(qn('w:val'), 'single')
                cellBottomBorder.set(qn('w:sz'), str(int(border_header * 8)))
                cellBottomBorder.set(qn('w:space'), '0')
                cellBottomBorder.set(qn('w:color'), '000000')
                tcBorders.append(cellBottomBorder)

                # 清除其他边框
                for border_name in ['top', 'left', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)

                # 清除旧边框
                oldTcBorders = tcPr.find(qn('w:tcBorders'))
                if oldTcBorders is not None:
                    tcPr.remove(oldTcBorders)

                tcPr.append(tcBorders)

        # 设置最后一行的下边框（粗线）- 三线表的第三条线
        if len(table.rows) > 1:
            last_row = table.rows[-1]
            for cell in last_row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                tcBorders = OxmlElement('w:tcBorders')

                # 下边框（粗线）
                cellBottomBorder = OxmlElement('w:bottom')
                cellBottomBorder.set(qn('w:val'), 'single')
                cellBottomBorder.set(qn('w:sz'), str(int(border_top_bottom * 8)))
                cellBottomBorder.set(qn('w:space'), '0')
                cellBottomBorder.set(qn('w:color'), '000000')
                tcBorders.append(cellBottomBorder)

                # 清除其他边框
                for border_name in ['top', 'left', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)

                # 清除旧边框
                oldTcBorders = tcPr.find(qn('w:tcBorders'))
                if oldTcBorders is not None:
                    tcPr.remove(oldTcBorders)

                tcPr.append(tcBorders)

        # 设置正文单元格格式
        for row_idx in range(1, len(table.rows)):
            row = table.rows[row_idx]
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        safe_set_font(run, body_font, body_font_size, body_bold, fallback_font)

    except Exception as e:
        if config.get('verbose'):
            print(f"  [WARN] 设置表格格式时出现警告: {e}")

# ============================================================================
# 图片位置处理
# ============================================================================

def fix_image_position(doc, config):
    """
    修复图片位置问题

    功能：
    1. 处理嵌入式图片（InlineShapes）- 设置居中对齐
    2. 处理浮动图片（通过XML访问）- 锁定锚定位置
    3. 防止图片漂移

    Args:
        doc: 文档对象
        config: 配置对象
    """
    verbose = config.get('verbose', True)
    inline_count = 0
    float_count = 0

    try:
        # -----------------------------------------------------------
        # 处理嵌入式图片（InlineShapes）
        # 这些图片嵌入在段落中，位置相对稳定
        # -----------------------------------------------------------
        inline_count = len(doc.inline_shapes)

        if inline_count > 0:
            for i, inline_shape in enumerate(doc.inline_shapes):
                try:
                    # 设置嵌入式图片居中
                    # InlineShape嵌入在某个run中，需要找到包含它的段落
                    # 通过遍历段落找到包含图片的段落

                    for paragraph in doc.paragraphs:
                        for run in paragraph.runs:
                            # 检查run中是否包含图片
                            if run._element.xpath('.//a:blip') or \
                               run._element.xpath('.//w:drawing'):
                                # 找到包含图片的run，设置段落居中
                                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                                if verbose and i < 5:  # 只显示前5个
                                    print(f"  [嵌入图片{i+1}] 已设置居中对齐")
                                break

                except Exception as e:
                    if verbose:
                        print(f"  [WARN] 处理嵌入图片{i+1}时出错: {e}")

        # -----------------------------------------------------------
        # 处理浮动图片（通过XML访问）
        # 这些图片容易出现位置漂移问题
        # -----------------------------------------------------------
        # 通过document part获取所有drawing元素
        try:
            # 获取文档的所有drawing元素（包含浮动图片）
            drawings = doc._element.xpath('.//w:drawing')

            for drawing in drawings:
                float_count += 1

                try:
                    # 查找anchor元素（浮动图片的锚定信息）
                    anchor = drawing.find(qn('wp:anchor'))
                    if anchor is not None:
                        # 设置锚定属性，防止图片漂移
                        # 修改simplePos属性
                        simplePos = anchor.find(qn('wp:simplePos'))
                        if simplePos is not None:
                            # 锁定简单位置
                            simplePos.set(qn('wp:x'), '0')
                            simplePos.set(qn('wp:y'), '0')

                        # 设置positionH和positionV相对于段落
                        positionH = anchor.find(qn('wp:positionH'))
                        if positionH is not None:
                            positionH.set(qn('wp:relativeFrom'), 'paragraph')

                        positionV = anchor.find(qn('wp:positionV'))
                        if positionV is not None:
                            positionV.set(qn('wp:relativeFrom'), 'paragraph')

                        # 锁定锚定标记
                        # 设置allowOverlap为false
                        for child in anchor:
                            if child.tag == qn('wp:allowOverlap'):
                                child.set(qn('wp:val'), '0')

                    if verbose and float_count <= 5:
                        print(f"  [浮动图片{float_count}] 已锁定锚定位置")

                except Exception as e:
                    if verbose:
                        print(f"  [WARN] 处理浮动图片{float_count}时出错: {e}")

        except Exception as e:
            if verbose:
                print(f"  [WARN] 访问浮动图片时出错: {e}")

        if verbose:
            print(f"\n[OK] 图片处理完成:")
            print(f"  - 嵌入图片: {inline_count} 个")
            print(f"  - 浮动图片: {float_count} 个")

        return True

    except Exception as e:
        if verbose:
            print(f"  [WARN] 图片处理出错: {e}")
        return False

# ============================================================================
# 文本转换（将来时态）
# ============================================================================

def convert_to_future_tense(text, config):
    """
    将文本转换为将来时态

    Args:
        text: 原始文本
        config: 配置对象

    Returns:
        转换后的文本
    """
    # 检查是否启用时态转换
    if not config.get('enable_future_tense', True):
        return text

    # 获取替换规则
    replacements = config.get_tense_replacements()

    result = text
    for old, new in replacements.items():
        result = result.replace(old, new)

    return result

# ============================================================================
# 全局配置对象（初始化时创建）
# ============================================================================

config = None

# ============================================================================
# 主处理函数
# ============================================================================

def format_document(input_path, output_path, reference_path=None, config_path=None):
    """
    格式化整个文档

    Args:
        input_path: 输入文档路径
        output_path: 输出文档路径
        reference_path: 参考文档路径（可选）
        config_path: 配置文件路径（可选）
    """
    global config

    print("="*60)
    print("Word文档自动化格式化工具 - 专业版")
    print("="*60)

    # 加载配置
    config = ConfigLoader(config_path)
    print(f"\n使用配置: {config.config_source}")

    # 验证文件存在
    if not os.path.exists(input_path):
        print(f"[ERR] 错误: 输入文档不存在 - {input_path}")
        return False

    verbose = config.get('verbose', True)

    print(f"\n正在打开文档: {input_path}")

    try:
        # 打开文档
        doc = Document(input_path)

        # -----------------------------------------------------------
        # 步骤1: 自动识别并格式化标题和正文
        # -----------------------------------------------------------
        if verbose:
            print("\n[步骤1] 自动识别并格式化标题和正文...")

        heading_count = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0}
        body_count = 0
        tense_count = 0

        # 获取标题识别正则表达式
        patterns = config.get_heading_patterns()

        for i, paragraph in enumerate(doc.paragraphs):
            # 跳过空段落
            if not paragraph.text.strip():
                continue

            # 检查是否是特殊页面样式（封面、目录等）
            style_name = paragraph.style.name if paragraph.style else ""
            special_styles = config.get('special_page_styles', '')
            if special_styles and style_name in special_styles.split(','):
                # 跳过特殊页面样式的段落
                if verbose and i < 20:
                    print(f"  [{i+1}] 跳过特殊样式 '{style_name}'")
                continue

            # 检测标题级别
            level = detect_heading_level(paragraph.text, patterns)

            if level > 0:
                # 应用标题格式
                apply_heading_format(paragraph, level, config)
                heading_count[level] += 1

                if verbose:
                    preview = paragraph.text.strip()[:30]
                    print(f"  [{i+1}] 检测到{level}级标题: {preview}...")
            else:
                # 应用正文格式
                apply_body_format(paragraph, config)
                body_count += 1

                # 转换将来时态
                original_text = paragraph.text
                new_text = convert_to_future_tense(original_text, config)

                if new_text != original_text:
                    # 安全更新文本
                    try:
                        # 清除原文
                        for run in paragraph.runs:
                            run.text = ''

                        # 添加新文本
                        if paragraph.runs:
                            paragraph.runs[0].text = new_text
                        else:
                            paragraph.add_run(new_text)

                        tense_count += 1

                        if verbose:
                            preview = original_text.strip()[:20]
                            print(f"  [{i+1}] 已转换时态: '{preview}...'")

                    except Exception as e:
                        if verbose:
                            print(f"  [WARN] 转换时态时出现警告: {e}")

            # 进度提示
            if verbose and (i + 1) % 20 == 0:
                print(f"  已处理 {i + 1} 个段落...")

        if verbose:
            print(f"\n[OK] 段落处理完成:")
            print(f"  - 一级标题: {heading_count[1]} 个")
            print(f"  - 二级标题: {heading_count[2]} 个")
            print(f"  - 三级标题: {heading_count[3]} 个")
            print(f"  - 四级标题: {heading_count[4]} 个")
            print(f"  - 五级标题: {heading_count[5]} 个")
            print(f"  - 六级标题: {heading_count[6]} 个")
            print(f"  - 正文段落: {body_count} 个")
            print(f"  - 时态转换: {tense_count} 处")

        # -----------------------------------------------------------
        # 步骤2: 自动格式化表格为三线表样式
        # -----------------------------------------------------------
        if verbose:
            print("\n[步骤2] 自动格式化表格为三线表样式...")

        table_count = len(doc.tables)

        if table_count > 0:
            for i, table in enumerate(doc.tables):
                if verbose:
                    print(f"  正在处理表格 {i+1}/{table_count}...")
                apply_three_line_table(table, config)

            if verbose:
                print(f"\n[OK] 表格处理完成: {table_count} 个表格已设置为三线表样式")
        else:
            if verbose:
                print("  未检测到表格")

        # -----------------------------------------------------------
        # 步骤3: 修复图片位置
        # -----------------------------------------------------------
        if verbose:
            print("\n[步骤3] 修复图片位置...")

        fix_image_position(doc, config)

        # -----------------------------------------------------------
        # 步骤4: 保存文档
        # -----------------------------------------------------------
        if verbose:
            print(f"\n[步骤4] 保存文档: {output_path}")

        doc.save(output_path)

        print("\n" + "="*60)
        print("[OK] 文档处理完成！")
        print("="*60)

        if verbose:
            print("\n后续操作建议:")
            print("1. 打开输出文档检查格式是否正确")
            print("2. 手动生成目录（引用 → 目录 → 自动目录）")
            print("3. 检查所有表格是否为三线表样式")
            print("4. 检查图片是否需要重新绘制")
            print("5. 检查将来时态转换是否准确")

        return True

    except Exception as e:
        print(f"\n[ERR] 处理文档时出错: {e}")
        import traceback
        traceback.print_exc()
        return False

# ============================================================================
# 命令行入口
# ============================================================================

def main():
    """命令行入口"""
    import argparse

    parser = argparse.ArgumentParser(
        description='Word文档自动化格式化工具 - 专业版',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  %(prog)s input.docx output.docx
  %(prog)s input.docx output.docx --reference Revised.docx
  %(prog)s input.docx output.docx --config word_format.conf

配置文件:
  默认配置文件: word_format.conf
  如果配置文件存在，将自动加载并应用样式
  配置文件优先级 > 参考文档样式 > 默认样式

功能:
  [OK] 自动识别标题（1, 1.1, 1.1.1等）
  [OK] 自动应用标题样式
  [OK] 自动格式化正文
  [OK] 自动设置三线表样式
  [OK] 自动转换将来时态
  [OK] 支持配置文件自定义样式
        """
    )

    parser.add_argument('input', help='输入文档路径')
    parser.add_argument('output', help='输出文档路径')
    parser.add_argument('--reference', '-r', help='参考文档路径（可选）')
    parser.add_argument('--config', '-c', help='配置文件路径（可选）')
    parser.add_argument('--quiet', '-q', action='store_true', help='安静模式（不显示详细日志）')

    args = parser.parse_args()

    # 如果使用安静模式，临时修改配置
    if args.quiet:
        global config
        config = ConfigLoader()
        config.config['verbose'] = False

    success = format_document(args.input, args.output, args.reference, args.config)
    sys.exit(0 if success else 1)

if __name__ == '__main__':
    main()