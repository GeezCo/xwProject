# -*- coding: utf-8 -*-
"""
读取Word文档样式信息工具
用于分析参考文档的样式配置
"""

from docx import Document
from docx.shared import Pt
import sys

def read_document_styles(doc_path):
    """
    读取文档样式信息

    Args:
        doc_path: 文档路径
    """
    print("="*60)
    print("Word文档样式分析工具")
    print("="*60)

    print(f"\n正在读取文档: {doc_path}")

    try:
        doc = Document(doc_path)

        # -----------------------------------------------------------
        # 1. 分析段落样式（前30个非空段落）
        # -----------------------------------------------------------
        print("\n[段落样式分析]")
        print("-"*60)

        paragraph_styles = {}
        sample_count = 0

        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue

            sample_count += 1
            if sample_count > 30:
                break

            # 获取段落样式信息
            style_name = para.style.name if para.style else "未知"
            text_preview = para.text.strip()[:40]

            # 获取字体信息
            font_info = {}
            if para.runs:
                run = para.runs[0]
                font_info['font_name'] = run.font.name
                font_info['font_size'] = run.font.size.pt if run.font.size else "默认"
                font_info['bold'] = run.font.bold

                # 获取东亚字体
                try:
                    rPr = run._element.rPr
                    if rPr:
                        rFonts = rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                        if rFonts:
                            east_asia = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                            if east_asia:
                                font_info['east_asia_font'] = east_asia
                except:
                    pass

            # 获取段落格式信息
            para_format = para.paragraph_format
            format_info = {}
            format_info['line_spacing'] = para_format.line_spacing.pt if para_format.line_spacing else "默认"
            format_info['first_line_indent'] = para_format.first_line_indent

            # 打印样式信息
            print(f"\n段落 {sample_count}: '{text_preview}...'")
            print(f"  样式名称: {style_name}")
            if font_info:
                print(f"  字体: {font_info.get('font_name', '默认')}")
                if 'east_asia_font' in font_info:
                    print(f"  中文字体: {font_info['east_asia_font']}")
                print(f"  字号: {font_info.get('font_size', '默认')} 磅")
                print(f"  加粗: {font_info.get('bold', False)}")
            if format_info:
                print(f"  行距: {format_info.get('line_spacing', '默认')} 磅")
                indent = format_info.get('first_line_indent')
                if indent:
                    indent_chars = indent.pt / 12 if indent.pt else 0
                    print(f"  首行缩进: {indent_chars:.1f} 字符")

            # 统计样式
            if style_name not in paragraph_styles:
                paragraph_styles[style_name] = {
                    'count': 0,
                    'font_info': font_info,
                    'format_info': format_info
                }
            paragraph_styles[style_name]['count'] += 1

        # -----------------------------------------------------------
        # 2. 样式统计汇总
        # -----------------------------------------------------------
        print("\n" + "="*60)
        print("[样式统计汇总]")
        print("-"*60)

        for style_name, info in paragraph_styles.items():
            print(f"\n样式: {style_name}")
            print(f"  出现次数: {info['count']}")
            if info['font_info']:
                fi = info['font_info']
                print(f"  字体配置:")
                print(f"    - 西文: {fi.get('font_name', '默认')}")
                if 'east_asia_font' in fi:
                    print(f"    - 中文: {fi['east_asia_font']}")
                print(f"    - 字号: {fi.get('font_size', '默认')} 磅")
                print(f"    - 加粗: {fi.get('bold', False)}")
            if info['format_info']:
                fi = info['format_info']
                print(f"  格式配置:")
                print(f"    - 行距: {fi.get('line_spacing', '默认')} 磅")

        # -----------------------------------------------------------
        # 3. 分析表格样式
        # -----------------------------------------------------------
        print("\n" + "="*60)
        print("[表格样式分析]")
        print("-"*60)

        table_count = len(doc.tables)
        print(f"\n表格总数: {table_count}")

        if table_count > 0:
            # 分析第一个表格
            table = doc.tables[0]
            print(f"\n表格1 分析:")
            print(f"  行数: {len(table.rows)}")
            print(f"  列数: {len(table.columns)}")

            # 尝试获取边框信息
            try:
                tbl = table._tbl
                tblPr = tbl.tblPr
                if tblPr:
                    tblBorders = tblPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
                    if tblBorders:
                        print("  边框配置:")
                        for border in tblBorders:
                            border_name = border.tag.split('}')[1] if '}' in border.tag else border.tag
                            val = border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            sz = border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz')
                            if sz:
                                sz_pt = int(sz) / 8  # 转换为磅
                                print(f"    - {border_name}: {val}, {sz_pt}磅")
            except Exception as e:
                print(f"  [无法获取边框信息: {e}]")

        # -----------------------------------------------------------
        # 4. 首页特殊布局检查
        # -----------------------------------------------------------
        print("\n" + "="*60)
        print("[首页特殊布局检查]")
        print("-"*60)

        # 检查前5个段落是否是特殊页面（封面等）
        print("\n检查前5个段落:")
        for i, para in enumerate(doc.paragraphs[:5]):
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else "未知"
                print(f"  段落{i+1}: '{text[:30]}...' - 样式: {style_name}")

        print("\n" + "="*60)
        print("分析完成！")
        print("="*60)

        return True

    except Exception as e:
        print(f"\n错误: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("使用方法: python read_styles.py <文档路径>")
        sys.exit(1)

    read_document_styles(sys.argv[1])