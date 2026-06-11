# -*- coding: utf-8 -*-
"""
只修改Word文档字体
中文：宋体
英文：Times New Roman
其他格式不变
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import sys

def change_font_only(input_path, output_path):
    """
    只修改文档字体，其他格式不变

    Args:
        input_path: 输入文档路径
        output_path: 输出文档路径
    """
    print("="*70)
    print("只修改字体 - 其他格式不变")
    print("="*70)
    print(f"\n输入文档: {input_path}")
    print(f"输出文档: {output_path}")

    # 打开文档
    doc = Document(input_path)

    # 统计
    total_paragraphs = 0
    changed_runs = 0

    print("\n正在修改字体...")

    # 遍历所有段落
    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        total_paragraphs += 1

        # 修改每个run的字体
        for run in para.runs:
            try:
                # 设置西文字体为Times New Roman
                run.font.name = 'Times New Roman'

                # 设置中文字体为宋体
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), '宋体')
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')

                changed_runs += 1
            except Exception as e:
                print(f"  [WARN] 修改字体失败: {e}")

    # 修改表格中的字体
    table_count = 0
    for table in doc.tables:
        table_count += 1
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        try:
                            run.font.name = 'Times New Roman'
                            rPr = run._element.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:eastAsia'), '宋体')
                            rFonts.set(qn('w:ascii'), 'Times New Roman')
                            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                            changed_runs += 1
                        except:
                            pass

    # 保存文档
    doc.save(output_path)

    print(f"\n[OK] 完成！")
    print(f"  - 处理段落: {total_paragraphs} 个")
    print(f"  - 处理表格: {table_count} 个")
    print(f"  - 修改文本块: {changed_runs} 个")
    print(f"  - 字体设置: 中文=宋体, 英文=Times New Roman")
    print("="*70)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法: python change_font_only.py <输入文档> <输出文档>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    change_font_only(input_file, output_file)
