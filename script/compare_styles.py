# -*- coding: utf-8 -*-
"""
详细读取甲方规格模板样式
逐段落对比样式信息
"""

from docx import Document
from docx.shared import Pt
import sys

def detailed_style_analysis(doc_path):
    """详细分析文档样式"""
    print("="*70)
    print("甲方规格模板样式详细分析")
    print("="*70)
    print(f"\n文档: {doc_path}")

    doc = Document(doc_path)

    # 分析前50个非空段落
    print("\n【段落样式详细分析】")
    print("-"*70)

    count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        count += 1
        if count > 50:
            break

        style_name = para.style.name if para.style else "未知"

        # 字体信息
        font_name = None
        font_size = None
        east_asia_font = None
        bold = None

        if para.runs:
            run = para.runs[0]
            font_name = run.font.name
            font_size = run.font.size.pt if run.font.size else None
            bold = run.font.bold

            # 东亚字体
            try:
                rPr = run._element.rPr
                if rPr is not None:
                    rFonts = rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                    if rFonts is not None:
                        east_asia = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                        if east_asia:
                            east_asia_font = east_asia
            except:
                pass

        # 段落格式
        pf = para.paragraph_format
        line_spacing = pf.line_spacing.pt if pf.line_spacing else None
        first_indent = pf.first_line_indent.pt if pf.first_line_indent else None
        left_indent = pf.left_indent.pt if pf.left_indent else None

        # 输出
        preview = text[:40] if len(text) > 40 else text
        print(f"\n[{count}] '{preview}...'")
        print(f"  样式名: {style_name}")
        if font_name:
            print(f"  西文字体: {font_name}")
        if east_asia_font:
            print(f"  中文字体: {east_asia_font}")
        if font_size:
            print(f"  字号: {font_size}磅")
        if bold:
            print(f"  加粗: {bold}")
        if line_spacing:
            print(f"  行距: {line_spacing}磅")
        if first_indent:
            print(f"  首行缩进: {first_indent}磅 ({first_indent/12:.1f}字符)")
        if left_indent:
            print(f"  左缩进: {left_indent}磅 ({left_indent/12:.1f}字符)")

    # 汇总关键样式
    print("\n" + "="*70)
    print("【关键样式汇总】")
    print("-"*70)

    # 找正文段落（从"1 范围"开始）
    print("\n正文样式（从1范围开始）:")
    found_body = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("1 ") and ("范围" in text or "引言" in text):
            found_body = True
            # 分析接下来的正文段落
            for j in range(i, min(i+10, len(doc.paragraphs))):
                p = doc.paragraphs[j]
                if p.text.strip() and not p.text.strip().startswith(("1", "2", "3", "4", "5", "6")):
                    if p.runs:
                        run = p.runs[0]
                        print(f"\n正文段落: '{p.text.strip()[:30]}...'")
                        print(f"  字体: {run.font.name}")
                        if run.font.size:
                            print(f"  字号: {run.font.size.pt}磅")
                        try:
                            rPr = run._element.rPr
                            if rPr:
                                rFonts = rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                                if rFonts:
                                    ea = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                                    if ea:
                                        print(f"  中文字体: {ea}")
                        except:
                            pass
                        pf = p.paragraph_format
                        if pf.line_spacing:
                            print(f"  行距: {pf.line_spacing.pt}磅")
                        if pf.first_line_indent:
                            print(f"  首行缩进: {pf.first_line_indent.pt}磅")
                        break
            break

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python compare_styles.py <文档>")
        sys.exit(1)

    detailed_style_analysis(sys.argv[1])