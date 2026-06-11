# -*- coding: utf-8 -*-
"""
完整读取甲方规格模板所有样式
逐段详细分析，不能遗漏
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import sys

def full_analysis(doc_path):
    """完整分析文档所有样式"""

    doc = Document(doc_path)

    print("="*80)
    print("甲方规格模板完整样式分析")
    print("="*80)

    # 分析所有段落
    total = len(doc.paragraphs)
    print(f"\n段落总数: {total}")

    print("\n【逐段样式分析】")
    print("-"*80)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else "未知"

        # 字体信息
        font_info = {}
        if para.runs:
            run = para.runs[0]
            font_info['name'] = run.font.name
            font_info['size'] = run.font.size.pt if run.font.size else None
            font_info['bold'] = run.font.bold

            # 东亚字体
            try:
                rPr = run._element.rPr
                if rPr:
                    rFonts = rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                    if rFonts:
                        ea = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                        ascii_f = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')
                        hansi = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi')
                        if ea: font_info['eastAsia'] = ea
                        if ascii_f: font_info['ascii'] = ascii_f
                        if hansi: font_info['hAnsi'] = hansi
            except:
                pass

        # 段落格式
        pf = para.paragraph_format
        para_info = {}
        if pf.line_spacing: para_info['line_spacing'] = pf.line_spacing.pt
        if pf.first_line_indent: para_info['first_indent'] = pf.first_line_indent.pt
        if pf.left_indent: para_info['left_indent'] = pf.left_indent.pt
        if pf.alignment: para_info['alignment'] = pf.alignment

        # 输出
        preview = text[:50] if len(text) > 50 else text

        # 确定字体显示
        font_display = ""
        if font_info.get('eastAsia'):
            font_display = font_info['eastAsia']
        elif font_info.get('hAnsi'):
            font_display = font_info['hAnsi']
        elif font_info.get('name'):
            font_display = font_info['name']

        print(f"\n[{i+1}] 样式:{style_name}")
        print(f"    内容: '{preview}'")
        if font_display:
            print(f"    字体: {font_display}")
        if font_info.get('size'):
            print(f"    字号: {font_info['size']}磅")
        if font_info.get('bold'):
            print(f"    加粗: {font_info['bold']}")
        if para_info.get('line_spacing'):
            print(f"    行距: {para_info['line_spacing']}磅")
        if para_info.get('first_indent'):
            indent_char = para_info['first_indent'] / 12
            print(f"    首行缩进: {para_info['first_indent']}磅 ({indent_char:.1f}字符)")
        if para_info.get('left_indent'):
            indent_char = para_info['left_indent'] / 12
            print(f"    左缩进: {para_info['left_indent']}磅 ({indent_char:.1f}字符)")

        # 只显示前100个段落
        if i >= 100:
            print("\n... 已显示前100个段落，后续省略 ...")
            break

    # 分析正文部分（从目录后开始）
    print("\n" + "="*80)
    print("【正文部分详细分析】")
    print("="*80)

    # 找正文开始位置（目录之后）
    body_start = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # 正文从"范围"开始（目录之后）
        if text in ["范围", "引言"] or text.startswith("1 ") and "范围" in text:
            body_start = i
            break

    print(f"\n正文起始位置: 段落{body_start+1}")

    # 分析正文前20个段落
    body_count = 0
    for i in range(body_start, min(body_start+30, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if not text:
            continue

        body_count += 1
        style_name = para.style.name if para.style else "未知"

        # 详细字体
        font_detail = ""
        if para.runs:
            run = para.runs[0]
            try:
                rPr = run._element.rPr
                if rPr:
                    rFonts = rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                    if rFonts:
                        ea = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                        ascii_f = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')
                        if ea: font_detail += f"中文:{ea} "
                        if ascii_f: font_detail += f"西文:{ascii_f} "
            except:
                pass

            size = run.font.size.pt if run.font.size else None
            if size: font_detail += f"字号:{size}磅 "

        print(f"\n正文段落{body_count}: '{text[:40]}'")
        print(f"  样式名: {style_name}")
        print(f"  字体: {font_detail}")

        # 行距缩进
        pf = para.paragraph_format
        if pf.line_spacing:
            print(f"  行距: {pf.line_spacing.pt}磅")
        if pf.first_line_indent:
            print(f"  首行缩进: {pf.first_line_indent.pt}磅")
        if pf.left_indent:
            print(f"  左缩进: {pf.left_indent.pt}磅")

    print("\n" + "="*80)
    print("分析完成")
    print("="*80)

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python full_analysis.py <文档>")
        sys.exit(1)

    full_analysis(sys.argv[1])