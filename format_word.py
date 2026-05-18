# -*- coding: utf-8 -*-
"""
Word文档格式化工具
自动应用软件设计说明文档的标准格式
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_paragraph_format(paragraph, font_name, font_size, line_spacing, first_line_indent=None):
    """
    设置段落格式

    Args:
        paragraph: 段落对象
        font_name: 字体名称
        font_size: 字体大小（磅）
        line_spacing: 行距（磅）
        first_line_indent: 首行缩进（字符数）
    """
    # 设置字体
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)

    # 设置段落格式
    paragraph.paragraph_format.line_spacing = Pt(line_spacing)

    # 设置首行缩进
    if first_line_indent:
        # 1个字符约等于12磅
        indent_pt = first_line_indent * 12
        paragraph.paragraph_format.first_line_indent = Pt(indent_pt)

def create_three_line_table(document):
    """
    创建三线格表格模板

    Returns:
        表格样式说明
    """
    # 说明如何创建三线格表格
    print("""
三线格表格创建步骤：
1. 插入表格
2. 选中表格，右键 → 表格属性 → 边框和底纹
3. 设置：
   - 上边框：1.5磅，黑色
   - 下边框：1.5磅，黑色
   - 表头行下边框：1磅，黑色
   - 其他边框：无
4. 表头单元格：
   - 字体：黑体，四号
   - 对齐：居中
5. 内容单元格：
   - 字体：宋体，四号
   - 对齐：左对齐
""")

def apply_heading_style(paragraph, level=1):
    """
    应用标题样式

    Args:
        paragraph: 段落对象
        level: 标题级别（1-6）
    """
    if level == 1:
        # 一级标题：黑体，三号，29磅行距
        set_paragraph_format(paragraph, '黑体', 16, 29)
    else:
        # 其他标题：仿宋GB2312，三号，29磅行距
        set_paragraph_format(paragraph, '仿宋GB2312', 16, 29)

def apply_body_style(paragraph):
    """
    应用正文样式
    正文：仿宋GB2312，三号，29磅行距，首行缩进2字符
    """
    set_paragraph_format(paragraph, '仿宋GB2312', 16, 29, first_line_indent=2)

def convert_to_future_tense(text):
    """
    将文本转换为将来时态

    Args:
        text: 原始文本

    Returns:
        转换后的文本
    """
    # 这里可以添加将来时态转换的逻辑
    # 示例替换规则
    replacements = {
        '已完成': '将完成',
        '已实现': '将实现',
        '正在进行': '将进行',
        '已经': '将要',
        '正在': '将要',
        '实现了': '将实现',
        '完成了': '将完成',
        '开发了': '将开发',
        '设计了': '将设计',
        '测试了': '将测试',
        '部署了': '将部署',
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    return text

def format_document(input_path, output_path):
    """
    格式化整个文档

    Args:
        input_path: 输入文档路径
        output_path: 输出文档路径
    """
    print(f"正在处理文档: {input_path}")

    try:
        # 打开文档
        doc = Document(input_path)

        # 处理段落
        for i, paragraph in enumerate(doc.paragraphs):
            # 跳过空段落
            if not paragraph.text.strip():
                continue

            # 判断段落类型
            if paragraph.style.name.startswith('Heading'):
                # 标题段落
                level = int(paragraph.style.name.replace('Heading ', ''))
                apply_heading_style(paragraph, level)
            else:
                # 正文段落
                apply_body_style(paragraph)

                # 转换为将来时态
                if paragraph.text.strip():
                    # 保留原文内容，只更新文本
                    text = convert_to_future_tense(paragraph.text)
                    if text != paragraph.text:
                        # 清除原文
                        for run in paragraph.runs:
                            run.text = ''
                        # 添加新文本
                        if paragraph.runs:
                            paragraph.runs[0].text = text
                        else:
                            paragraph.add_run(text)

            # 进度提示
            if (i + 1) % 10 == 0:
                print(f"已处理 {i + 1} 个段落...")

        # 处理表格
        for table in doc.tables:
            # 这里可以添加表格格式化逻辑
            print("检测到表格，请手动设置为三线格格式")

        # 保存文档
        doc.save(output_path)
        print(f"\n✓ 文档处理完成: {output_path}")
        print(f"✓ 共处理 {len(doc.paragraphs)} 个段落")

        # 输出后续提示
        print("\n后续手动操作：")
        print("1. 检查所有表格，应用三线格格式")
        print("2. 检查所有图片，重新绘制线框图")
        print("3. 检查格式是否正确")
        print("4. 检查时态转换是否准确")

    except Exception as e:
        print(f"处理文档时出错: {e}")
        import traceback
        traceback.print_exc()

def main():
    """主函数"""
    import sys

    if len(sys.argv) < 3:
        print("使用方法:")
        print("  python format_word.py <输入文档> <输出文档>")
        print("\n示例:")
        print("  python format_word.py input.docx output.docx")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    format_document(input_path, output_path)

if __name__ == '__main__':
    main()