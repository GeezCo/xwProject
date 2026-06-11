#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown → DOCX 转换脚本（GJB 438B-2009 格式）

用法:
    python md_to_docx.py [输入.md] [输出.docx]

默认:
    python md_to_docx.py chapters_软件设计/软件设计说明.md 软件设计说明.docx

依赖:
    pip install python-docx Pillow
"""

import sys
import re
import os
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

try:
    from PIL import Image, ImageDraw, ImageFont
    HAS_PIL = True
except ImportError:
    HAS_PIL = False


# ============================================================
# 字体和格式常量（GJB 438B-2009）
# ============================================================
FONT_HEITI = '黑体'
FONT_FANGSONG = '仿宋'
FONT_SONGTI = '宋体'
FONT_CONSOLAS = 'Consolas'

# 磅值
PT_H1 = 22        # 2号 = 22pt
PT_H2345 = 16     # 3号 = 16pt
PT_BODY = 16       # 3号 = 16pt
PT_TABLE = 14      # 4号 = 14pt
PT_CAPTION = 12    # 小四 = 12pt

# 行距（磅）
LINE_SPACING_28 = Pt(28)
LINE_SPACING_25 = Pt(25)

# 首行缩进
FIRST_LINE_INDENT = Cm(0.85)  # 约2个中文字符宽度（16pt字体下约0.85cm）


def set_run_font(run, font_name, font_size, bold=False, color=None):
    """设置 run 的字体属性"""
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = font_name
    # 设置中文字体
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    if color:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, line_spacing=None, space_before=None, space_after=None):
    """设置段落行距和段间距"""
    pf = paragraph.paragraph_format
    if line_spacing:
        pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after


def set_first_line_indent(paragraph, indent=FIRST_LINE_INDENT):
    """设置首行缩进"""
    paragraph.paragraph_format.first_line_indent = indent


def add_heading_paragraph(doc, text, level):
    """添加标题段落（不使用 Word 内置 Heading 样式，直接格式化）"""
    p = doc.add_paragraph()
    run = p.add_run(text)

    if level == 1:
        set_run_font(run, FONT_HEITI, PT_H1, bold=True)
    else:
        set_run_font(run, FONT_FANGSONG, PT_H2345, bold=True)

    set_paragraph_spacing(p, LINE_SPACING_28, space_before=Pt(12), space_after=Pt(6))
    return p


def add_body_paragraph(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph()
    # 处理全角空格缩进：如果以全角空格开头则不再额外添加首行缩进
    stripped = text.lstrip('　')
    if text.startswith('　'):
        set_first_line_indent(p, FIRST_LINE_INDENT)
        run = p.add_run(stripped)
    else:
        set_first_line_indent(p, FIRST_LINE_INDENT)
        run = p.add_run(text)

    set_run_font(run, FONT_FANGSONG, PT_BODY)
    set_paragraph_spacing(p, LINE_SPACING_28)
    return p


def add_caption_paragraph(doc, text):
    """添加表例或图例段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, FONT_HEITI, PT_CAPTION, bold=True)
    set_paragraph_spacing(p, line_spacing=Pt(20), space_before=Pt(6), space_after=Pt(3))
    return p


def add_list_paragraph(doc, text):
    """添加分项列表段落（a) b) c) 或 1) 2) 3) 格式）"""
    p = doc.add_paragraph()
    # 根据缩进层级设置左缩进
    indent_count = 0
    original_text = text
    while text.startswith('   '):
        text = text[3:]
        indent_count += 1

    p.paragraph_format.left_indent = Cm(0.85 * (indent_count + 1))
    p.paragraph_format.first_line_indent = Cm(0)

    run = p.add_run(text.strip())
    set_run_font(run, FONT_FANGSONG, PT_BODY)
    set_paragraph_spacing(p, LINE_SPACING_28)
    return p


def create_table(doc, headers, rows):
    """创建格式化的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # 表头行
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, FONT_HEITI, PT_TABLE, bold=True)
        set_paragraph_spacing(p, LINE_SPACING_25)
        # 表头背景色
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            set_run_font(run, FONT_SONGTI, PT_TABLE)
            set_paragraph_spacing(p, LINE_SPACING_25)

    return table


def ascii_art_to_image(text, width=800):
    """将 ASCII 线框图转为 PNG 图片"""
    if not HAS_PIL:
        return None

    lines = text.split('\n')
    font_size = 14
    try:
        font = ImageFont.truetype('consola.ttf', font_size)
    except (OSError, IOError):
        try:
            font = ImageFont.truetype('DejaVuSansMono.ttf', font_size)
        except (OSError, IOError):
            font = ImageFont.load_default()

    char_width = font_size * 0.6
    line_height = font_size * 1.4
    max_line_len = max(len(line) for line in lines) if lines else 0

    img_width = int(max_line_len * char_width) + 40
    img_height = int(len(lines) * line_height) + 40

    img = Image.new('RGB', (img_width, img_height), 'white')
    draw = ImageDraw.Draw(img)

    y = 20
    for line in lines:
        draw.text((20, y), line, fill='black', font=font)
        y += line_height

    buf = BytesIO()
    img.save(buf, format='PNG', dpi=(300, 300))
    buf.seek(0)
    return buf


def add_cover_page(doc):
    """添加封面页"""
    # 空行留白
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, Pt(28))

    # 密级
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('密级：内部')
    set_run_font(run, FONT_HEITI, 14)

    # 空行
    for _ in range(2):
        doc.add_paragraph()

    # 文档标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('献微系统技术文件')
    set_run_font(run, FONT_HEITI, 26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('软件设计说明')
    set_run_font(run, FONT_HEITI, 22, bold=True)

    # 空行
    for _ in range(3):
        doc.add_paragraph()

    # 文档信息
    info_lines = [
        '文档标识号：XW-SDD-001',
        '版 本 号：V1.0.0',
        '',
        '编制单位：献微系统开发团队',
        '编制日期：2026年5月',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_run_font(run, FONT_FANGSONG, 16)
        set_paragraph_spacing(p, Pt(28))

    # 分页
    doc.add_page_break()


def add_modification_record(doc):
    """添加文档修改记录页"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('文档修改记录')
    set_run_font(run, FONT_HEITI, PT_H1, bold=True)
    set_paragraph_spacing(p, LINE_SPACING_28)

    add_caption_paragraph(doc, '表0-1 文档修改记录表')

    headers = ['版本号', '修改原因', '修改内容', '修改人', '日期', '备注']
    rows = [
        ['V1.0.0', '新建文档', '创建软件设计说明', '献微系统开发团队', '2026-05', '初始版本'],
    ]
    create_table(doc, headers, rows)

    doc.add_page_break()


def add_toc(doc):
    """添加目次页（通过 Word 域代码，打开后按 F9 更新）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目  次')
    set_run_font(run, FONT_HEITI, PT_H1, bold=True)
    set_paragraph_spacing(p, LINE_SPACING_28)

    # 插入 TOC 域代码
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-5" \\h \\z \\u </w:instrText>')
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar2)

    run4 = p.add_run('（请右键更新域以生成目录）')
    set_run_font(run4, FONT_FANGSONG, 12)

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar3)

    doc.add_page_break()


# ============================================================
# Markdown 解析器
# ============================================================

# 正则模式
RE_TABLE_CAPTION = re.compile(r'^表\d+-\d+\s+.+')
RE_FIGURE_CAPTION = re.compile(r'^图\d+-\d+\s+.+')
RE_LIST_ITEM = re.compile(r'^(\s*)(a\)|b\)|c\)|d\)|e\)|f\)|g\)|h\)|\d+\)|\d+\.\s|①|②|③|④|⑤|⑥|⑦|⑧|⑨|⑩)')
RE_TABLE_ROW = re.compile(r'^\|(.+)\|$')
RE_TABLE_SEP = re.compile(r'^\|[\s\-:|]+\|$')
RE_CODE_BLOCK = re.compile(r'^```')
RE_HEADING = re.compile(r'^(\d+(?:\.\d+)*)\s+(.+)$')


def detect_heading_level(text):
    """检测标题级别（基于数字编号格式：1, 1.1, 1.1.1 等）"""
    m = RE_HEADING.match(text)
    if m:
        num_part = m.group(1)
        title_text = m.group(2)
        level = num_part.count('.') + 1
        return level, text
    return None, text


def parse_table_block(lines, start_idx):
    """解析 Markdown 表格块，返回 (headers, rows, end_idx)"""
    headers = []
    rows = []
    idx = start_idx

    # 第一行是表头
    m = RE_TABLE_ROW.match(lines[idx])
    if m:
        headers = [c.strip() for c in m.group(1).split('|')]
        idx += 1

    # 第二行是分隔线
    if idx < len(lines) and RE_TABLE_SEP.match(lines[idx]):
        idx += 1

    # 数据行
    while idx < len(lines) and RE_TABLE_ROW.match(lines[idx]):
        m = RE_TABLE_ROW.match(lines[idx])
        cells = [c.strip() for c in m.group(1).split('|')]
        # 补齐列数
        while len(cells) < len(headers):
            cells.append('')
        rows.append(cells[:len(headers)])
        idx += 1

    return headers, rows, idx


def is_ascii_art(text):
    """判断代码块内容是否为 ASCII 线框图"""
    art_chars = set('┌┐└┘├┤┬┴┼─│▶▼↑↓←→')
    count = sum(1 for c in text if c in art_chars)
    return count > 5


def parse_and_convert(md_text, doc):
    """解析 Markdown 文本并转换为 DOCX 内容"""
    lines = md_text.split('\n')
    idx = 0
    in_code_block = False
    code_block_lines = []

    # 跳过封面、修改记录和目次的 Markdown 文本
    # （这些已通过专门函数生成，不需要从 MD 解析）
    skip_header = True

    while idx < len(lines):
        line = lines[idx]

        # 跳过文件头部的封面/修改记录/目次内容
        if skip_header:
            if RE_HEADING.match(line.strip()):
                # 遇到第一个数字编号标题（如 "1 范围"），开始正式解析
                skip_header = False
            else:
                idx += 1
                continue

        stripped = line.strip()

        # 代码块处理
        if RE_CODE_BLOCK.match(stripped):
            if in_code_block:
                # 代码块结束
                in_code_block = False
                code_text = '\n'.join(code_block_lines)

                if is_ascii_art(code_text):
                    # ASCII 线框图 → 转为图片
                    img_buf = ascii_art_to_image(code_text)
                    if img_buf:
                        doc.add_picture(img_buf, width=Inches(5.5))
                        last_p = doc.paragraphs[-1]
                        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        # 无法转图片时作为代码块插入
                        for code_line in code_block_lines:
                            p = doc.add_paragraph()
                            run = p.add_run(code_line)
                            set_run_font(run, FONT_CONSOLAS, 10)
                            p.paragraph_format.first_line_indent = Cm(0)
                else:
                    # 普通代码块
                    for code_line in code_block_lines:
                        p = doc.add_paragraph()
                        run = p.add_run(code_line)
                        set_run_font(run, FONT_CONSOLAS, 10)
                        p.paragraph_format.first_line_indent = Cm(0)
                        set_paragraph_spacing(p, Pt(16))

                code_block_lines = []
            else:
                in_code_block = True
                code_block_lines = []
            idx += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            idx += 1
            continue

        # 空行
        if not stripped:
            idx += 1
            continue

        # 分割线（跳过）
        if stripped.startswith('---'):
            idx += 1
            continue

        # 表例（表X-Y 格式）
        if RE_TABLE_CAPTION.match(stripped):
            add_caption_paragraph(doc, stripped)
            idx += 1
            continue

        # 图例（图X-Y 格式）
        if RE_FIGURE_CAPTION.match(stripped):
            add_caption_paragraph(doc, stripped)
            idx += 1
            continue

        # Markdown 表格
        if RE_TABLE_ROW.match(stripped):
            headers, rows, new_idx = parse_table_block(lines, idx)
            if headers and rows:
                create_table(doc, headers, rows)
            idx = new_idx
            continue

        # 标题检测
        level, title_text = detect_heading_level(stripped)
        if level is not None and level <= 5:
            add_heading_paragraph(doc, title_text, level)
            idx += 1
            continue

        # 分项列表（a) b) c) / 1) 2) 3) / ①②③）
        if RE_LIST_ITEM.match(stripped) or RE_LIST_ITEM.match(line):
            add_list_paragraph(doc, line)
            idx += 1
            continue

        # 普通正文
        add_body_paragraph(doc, stripped)
        idx += 1


def setup_page(doc):
    """设置页面格式（A4）"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def main():
    # 参数解析
    if len(sys.argv) >= 3:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
    elif len(sys.argv) == 2:
        input_file = sys.argv[1]
        output_file = os.path.splitext(input_file)[0] + '.docx'
    else:
        input_file = 'chapters_软件设计/软件设计说明.md'
        output_file = '软件设计说明.docx'

    print(f'输入: {input_file}')
    print(f'输出: {output_file}')

    # 读取 Markdown 文件
    encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
    md_text = None
    for enc in encodings:
        try:
            with open(input_file, 'r', encoding=enc) as f:
                md_text = f.read()
            break
        except (UnicodeDecodeError, UnicodeError):
            continue

    if md_text is None:
        print(f'错误: 无法读取文件 {input_file}')
        sys.exit(1)

    print(f'读取成功: {len(md_text)} 字符, {md_text.count(chr(10))} 行')

    # 创建 Word 文档
    doc = Document()
    setup_page(doc)

    # 添加封面、修改记录和目次
    add_cover_page(doc)
    add_modification_record(doc)
    add_toc(doc)

    # 解析 Markdown 并转换
    parse_and_convert(md_text, doc)

    # 保存
    doc.save(output_file)
    print(f'转换完成: {output_file}')
    print(f'文件大小: {os.path.getsize(output_file) / 1024:.1f} KB')


if __name__ == '__main__':
    main()
